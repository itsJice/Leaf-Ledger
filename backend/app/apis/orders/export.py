"""Render a grouped-by-vendor order view into a downloadable PO document.

Three formats, each with the fields the buyer needs on a purchase order:
picture · product name · vendor SKU · size · quantity · unit price · line total,
plus per-vendor subtotals and a grand total. Product images are downloaded once
(cached) and embedded as small thumbnails; a failed/blocked image is simply
skipped so the export never breaks.
"""
from __future__ import annotations

import io
from typing import Optional

import requests
from PIL import Image

_IMG_CACHE: dict[str, Optional[bytes]] = {}
_UA = {"User-Agent": "Mozilla/5.0 (Macintosh; Intel Mac OS X 10_15_7) AppleWebKit/537.36"}


def _thumb(url: Optional[str], box: int = 96) -> Optional[bytes]:
    """Download an image and return PNG thumbnail bytes (cached; None on failure)."""
    if not url:
        return None
    if url in _IMG_CACHE:
        return _IMG_CACHE[url]
    data: Optional[bytes] = None
    try:
        r = requests.get(url, timeout=6, headers=_UA)
        if r.ok and r.content:
            im = Image.open(io.BytesIO(r.content))
            if im.mode not in ("RGB", "RGBA"):
                im = im.convert("RGB")
            im.thumbnail((box, box))
            out = io.BytesIO()
            im.convert("RGB").save(out, format="PNG")
            data = out.getvalue()
    except Exception:
        data = None
    _IMG_CACHE[url] = data
    return data


def _money(n) -> str:
    return "" if n is None else f"${float(n):,.2f}"


def _date(view) -> str:
    d = view.get("created_at")
    try:
        return d.strftime("%b %d, %Y") if d else ""
    except Exception:
        return str(d or "")


def render(view: dict, fmt: str) -> tuple[bytes, str, str]:
    """Return (bytes, media_type, extension) for the requested format."""
    if fmt == "xlsx":
        return _xlsx(view), "application/vnd.openxmlformats-officedocument.spreadsheetml.sheet", "xlsx"
    if fmt == "docx":
        return _docx(view), "application/vnd.openxmlformats-officedocument.wordprocessingml.document", "docx"
    return _pdf(view), "application/pdf", "pdf"


# ── Excel ───────────────────────────────────────────────────────────────────
def _xlsx(view: dict) -> bytes:
    from openpyxl import Workbook
    from openpyxl.drawing.image import Image as XLImage
    from openpyxl.styles import Font, Alignment, PatternFill
    from openpyxl.utils import get_column_letter

    wb = Workbook()
    wb.remove(wb.active)
    header_fill = PatternFill("solid", fgColor="1F5D4C")
    header_font = Font(bold=True, color="FFFFFF")
    bold = Font(bold=True)
    cols = ["Image", "Product", "SKU", "Size", "Qty", "Unit", "Line Total", "Link"]
    widths = [14, 46, 16, 10, 6, 12, 12, 30]

    used_names: set[str] = set()
    for v in view["vendors"] or [{"supplier_name": "Order", "items": [], "subtotal": 0}]:
        title = "".join(c for c in v["supplier_name"] if c not in "[]:*?/\\")[:28] or "Vendor"
        base, n = title, 2
        while title in used_names:
            title = f"{base[:26]}_{n}"; n += 1
        used_names.add(title)
        ws = wb.create_sheet(title)
        ws.append([view["name"]])
        ws["A1"].font = Font(bold=True, size=14)
        ws.append([f"{v['supplier_name']}  ·  {_date(view)}"])
        ws.append([])
        ws.append(cols)
        hdr_row = ws.max_row
        for i, _c in enumerate(cols, start=1):
            cell = ws.cell(row=hdr_row, column=i)
            cell.fill = header_fill; cell.font = header_font
            ws.column_dimensions[get_column_letter(i)].width = widths[i - 1]

        for it in v["items"]:
            ws.append([
                "", it["name"], it["sku"] or "", it["size"] or "", it["quantity"],
                float(it["unit_price"]) if it["unit_price"] is not None else "",
                float(it["line_total"]) if it["line_total"] is not None else "",
                it["product_url"] or "",
            ])
            row = ws.max_row
            ws.row_dimensions[row].height = 60
            ws.cell(row=row, column=5).alignment = Alignment(horizontal="center")
            for c in (6, 7):
                ws.cell(row=row, column=c).number_format = '"$"#,##0.00'
            if it["product_url"]:
                link = ws.cell(row=row, column=8)
                link.hyperlink = it["product_url"]; link.value = "View on site"
                link.font = Font(color="1F5D4C", underline="single")
            png = _thumb(it["image_url"])
            if png:
                img = XLImage(io.BytesIO(png))
                img.width, img.height = 70, 70
                ws.add_image(img, f"A{row}")

        ws.append([])
        total_row = ws.max_row + 1
        ws.cell(row=total_row, column=4, value="Subtotal").font = bold
        ws.cell(row=total_row, column=5, value=v["subtotal_qty"]).font = bold
        st = ws.cell(row=total_row, column=7, value=float(v["subtotal"])); st.font = bold
        st.number_format = '"$"#,##0.00'

    out = io.BytesIO()
    wb.save(out)
    return out.getvalue()


# ── Word ────────────────────────────────────────────────────────────────────
def _docx(view: dict) -> bytes:
    from docx import Document
    from docx.shared import Pt, Inches, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH

    doc = Document()
    h = doc.add_heading(view["name"], level=0)
    sub = doc.add_paragraph(f"Purchase Order  ·  {_date(view)}")
    sub.runs[0].font.size = Pt(10)

    for v in view["vendors"]:
        doc.add_heading(v["supplier_name"], level=1)
        table = doc.add_table(rows=1, cols=7)
        table.style = "Light Grid Accent 1"
        hdr = table.rows[0].cells
        for i, label in enumerate(["Image", "Product", "SKU", "Size", "Qty", "Unit", "Total"]):
            hdr[i].text = label
            for p in hdr[i].paragraphs:
                for r in p.runs:
                    r.font.bold = True; r.font.size = Pt(9)
        for it in v["items"]:
            cells = table.add_row().cells
            png = _thumb(it["image_url"])
            if png:
                try:
                    cells[0].paragraphs[0].add_run().add_picture(io.BytesIO(png), width=Inches(0.7))
                except Exception:
                    pass
            cells[1].text = it["name"] or ""
            if it["product_url"]:
                pu = cells[1].add_paragraph().add_run(it["product_url"])
                pu.font.size = Pt(7); pu.font.color.rgb = RGBColor(0x1F, 0x5D, 0x4C)
            cells[2].text = it["sku"] or ""
            cells[3].text = it["size"] or ""
            cells[4].text = str(it["quantity"])
            cells[5].text = _money(it["unit_price"])
            cells[6].text = _money(it["line_total"])
            for i in (2, 3, 4, 5, 6):
                for p in cells[i].paragraphs:
                    for r in p.runs:
                        r.font.size = Pt(9)
        sp = doc.add_paragraph()
        run = sp.add_run(f"{v['supplier_name']} subtotal: {v['subtotal_qty']} items  —  {_money(v['subtotal'])}")
        run.font.bold = True; run.font.size = Pt(10)
        sp.alignment = WD_ALIGN_PARAGRAPH.RIGHT

    tot = doc.add_paragraph()
    tr = tot.add_run(f"Order total: {_money(view['total_cost'])}  ({view['total_qty']} items)")
    tr.font.bold = True; tr.font.size = Pt(12)
    tot.alignment = WD_ALIGN_PARAGRAPH.RIGHT

    out = io.BytesIO()
    doc.save(out)
    return out.getvalue()


# ── PDF ─────────────────────────────────────────────────────────────────────
def _pdf(view: dict) -> bytes:
    from reportlab.lib import colors
    from reportlab.lib.pagesizes import letter
    from reportlab.lib.units import inch
    from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
    from reportlab.platypus import (
        SimpleDocTemplate, Table, TableStyle, Paragraph, Spacer, Image as RLImage,
    )

    styles = getSampleStyleSheet()
    cell = ParagraphStyle("cell", parent=styles["Normal"], fontSize=8, leading=10)
    link = ParagraphStyle("link", parent=cell, textColor=colors.HexColor("#1F5D4C"), fontSize=7)
    accent = colors.HexColor("#1F5D4C")

    buf = io.BytesIO()
    doc = SimpleDocTemplate(buf, pagesize=letter, topMargin=0.6 * inch, bottomMargin=0.6 * inch,
                            leftMargin=0.5 * inch, rightMargin=0.5 * inch)
    story = [
        Paragraph(f"<b>{_esc(view['name'])}</b>", styles["Title"]),
        Paragraph(f"Purchase Order &nbsp;·&nbsp; {_date(view)}", styles["Normal"]),
        Spacer(1, 12),
    ]
    header = ["", "Product", "SKU", "Size", "Qty", "Unit", "Total"]
    col_widths = [0.85 * inch, 3.0 * inch, 1.0 * inch, 0.6 * inch, 0.45 * inch, 0.7 * inch, 0.8 * inch]

    for v in view["vendors"]:
        story.append(Paragraph(f"<b>{_esc(v['supplier_name'])}</b>", styles["Heading2"]))
        rows = [header]
        for it in v["items"]:
            png = _thumb(it["image_url"])
            img_flow = RLImage(io.BytesIO(png), width=0.7 * inch, height=0.7 * inch) if png else Paragraph("", cell)
            name = _esc(it["name"] or "")
            if it["product_url"]:
                name += f'<br/><font size=6><a href="{_esc(it["product_url"])}">{_esc(it["product_url"])[:48]}</a></font>'
            rows.append([
                img_flow, Paragraph(name, cell), Paragraph(_esc(it["sku"] or ""), cell),
                Paragraph(_esc(it["size"] or ""), cell), Paragraph(str(it["quantity"]), cell),
                Paragraph(_money(it["unit_price"]), cell), Paragraph(_money(it["line_total"]), cell),
            ])
        rows.append(["", "", "", "", str(v["subtotal_qty"]), Paragraph("<b>Subtotal</b>", cell),
                     Paragraph(f"<b>{_money(v['subtotal'])}</b>", cell)])
        t = Table(rows, colWidths=col_widths, repeatRows=1)
        t.setStyle(TableStyle([
            ("BACKGROUND", (0, 0), (-1, 0), accent),
            ("TEXTCOLOR", (0, 0), (-1, 0), colors.white),
            ("FONTSIZE", (0, 0), (-1, 0), 8),
            ("GRID", (0, 0), (-1, -1), 0.4, colors.HexColor("#DDD8D0")),
            ("VALIGN", (0, 0), (-1, -1), "MIDDLE"),
            ("ALIGN", (4, 0), (4, -1), "CENTER"),
            ("ALIGN", (5, 0), (6, -1), "RIGHT"),
            ("ROWBACKGROUNDS", (0, 1), (-1, -2), [colors.white, colors.HexColor("#FAF8F4")]),
        ]))
        story.append(t)
        story.append(Spacer(1, 16))

    story.append(Paragraph(
        f"<para align=right><b>Order total: {_money(view['total_cost'])}</b> "
        f"({view['total_qty']} items)</para>", styles["Normal"]))
    doc.build(story)
    return buf.getvalue()


def _esc(s: str) -> str:
    return (str(s or "").replace("&", "&amp;").replace("<", "&lt;").replace(">", "&gt;"))
